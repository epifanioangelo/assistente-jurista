#!/usr/bin/env python3
"""
Cria o template DOCX profissional para Agravo de Instrumento.

Layout gerado:
  - Cabeçalho: logo [AE] centralizado + informações de contato à direita
  - Rodapé: barra navy com nome do escritório + número de página dinâmico
  - Margens: top 4.5 cm (acomoda cabeçalho), bottom 2.5 cm, left 4 cm, right 2 cm
  - Barra lateral esquerda: texto vertical via caixa de texto XML

Saída: Modelos jus/AGRAVO_INSTRUMENTO.docx

Uso:
    python3 criar_template_agravo.py
"""

import io
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

ROOT    = Path(__file__).parent.parent
ASSETS  = ROOT / "_assets"
MODELOS = ROOT / "Modelos jus"

NAVY_HEX  = "0D2340"
NAVY2_HEX = "1A3A5C"
NAVY      = (13, 35, 64)
NAVY2     = (26, 58, 92)
WHITE     = (255, 255, 255)

TNR_B = "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"
HEL   = "/System/Library/Fonts/HelveticaNeue.ttc"

CONTACT = {
    "oab":    "OAB/SP XXXXX",
    "site":   "www.angeloepifanio.com.br",
    "email":  "contato@angeloepifanio.com.br",
    "tel":    "(12) 0000-0000",
    "cidade": "São José dos Campos/SP",
}


# ─── Image helpers ────────────────────────────────────────────────────────────

def _sw(d, txt, font, sp):
    return sum((d.textbbox((0,0),ch,font=font)[2]-d.textbbox((0,0),ch,font=font)[0])+sp
               for ch in txt) - sp

def _ds(d, txt, font, x, y, color, sp):
    cx = x
    for ch in txt:
        b = d.textbbox((0,0), ch, font=font)
        d.text((cx-b[0], y-b[1]), ch, font=font, fill=color)
        cx += (b[2]-b[0]) + sp


def gerar_header_png(info: dict = CONTACT) -> bytes:
    """PNG do cabeçalho (1700 × 260 px, 200 dpi)."""
    W, H, SC = 1700, 260, 3
    sw, sh = W*SC, H*SC
    s = SC
    img = Image.new("RGBA", (sw, sh), WHITE+(255,))
    d   = ImageDraw.Draw(img)

    f_ae  = ImageFont.truetype(TNR_B, 60*s)
    f_nm  = ImageFont.truetype(TNR_B, 14*s)
    f_adv = ImageFont.truetype(HEL,    9*s, index=0)
    f_ct  = ImageFont.truetype(HEL,    8*s, index=0)

    LCX = int(sw * 0.37)      # centro horizontal do bloco do logo
    lth = max(1, s//2)         # espessura das linhas decorativas

    # AE
    bae = d.textbbox((0,0), "AE", font=f_ae)
    ae_w, ae_h = bae[2]-bae[0], bae[3]-bae[1]
    ae_x = LCX - ae_w//2 - bae[0]
    ae_y = 14*s
    d.text((ae_x, ae_y), "AE", font=f_ae, fill=NAVY)

    # "ANGELO EPIFANIO" espaçado
    SP = int(2.8*s)
    nw  = _sw(d, "ANGELO EPIFANIO", f_nm, SP)
    nh  = d.textbbox((0,0),"A",font=f_nm)[3]-d.textbbox((0,0),"A",font=f_nm)[1]
    ny  = ae_y + ae_h + 6*s
    _ds(d, "ANGELO EPIFANIO", f_nm, LCX-nw//2, ny, NAVY, SP)

    # Linhas flanqueando o nome
    ly  = ny + nh + 6*s
    gap = 32*s
    ll  = 80*s
    d.rectangle([LCX-gap-ll, ly, LCX-gap, ly+lth], fill=NAVY2)
    d.rectangle([LCX+gap,    ly, LCX+gap+ll, ly+lth], fill=NAVY2)

    # ♦
    db  = d.textbbox((0,0),"♦",font=f_adv)
    dh  = db[3]-db[1]
    dyy = ly + 7*s
    _ds(d, "♦", f_adv, LCX-(db[2]-db[0])//2, dyy, NAVY, 0)

    # "A D V O C A C I A"
    SA  = int(2*s)
    aw  = _sw(d, "A D V O C A C I A", f_adv, SA)
    _ds(d, "A D V O C A C I A", f_adv, LCX-aw//2, dyy+dh+5*s, NAVY2, SA)

    # Contato (lado direito, alinhado à direita)
    rx = sw - 28*s
    cy = 20*s
    b  = d.textbbox((0,0), info["oab"], font=f_nm)
    d.text((rx-(b[2]-b[0])-b[0], cy-b[1]), info["oab"], font=f_nm, fill=NAVY)
    cy += (b[3]-b[1]) + 10*s
    d.rectangle([rx-130*s, cy, rx, cy+lth], fill=NAVY2)
    cy += 12*s
    for txt in [info["site"], info["email"], info["tel"]]:
        b = d.textbbox((0,0), txt, font=f_ct)
        d.text((rx-(b[2]-b[0])-b[0], cy-b[1]), txt, font=f_ct, fill=NAVY2)
        cy += (b[3]-b[1]) + 7*s

    img = img.resize((W, H), Image.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="PNG", dpi=(200, 200))
    return buf.getvalue()


# ─── XML helpers ─────────────────────────────────────────────────────────────

def set_cell_color(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _get_or_add_tblPr(tbl):
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    return tblPr


def remove_table_borders(table):
    tbl   = table._tbl
    tblPr = _get_or_add_tblPr(tbl)
    bdr   = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        bdr.append(el)
    tblPr.append(bdr)


def set_table_full_width(table):
    tbl   = table._tbl
    tblPr = _get_or_add_tblPr(tbl)
    tblW  = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)


def add_page_field(run):
    """Adiciona campo de número de página dinâmico ao run."""
    fld_b = OxmlElement("w:fldChar"); fld_b.set(qn("w:fldCharType"), "begin")
    ins   = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve"); ins.text = " PAGE "
    fld_s = OxmlElement("w:fldChar"); fld_s.set(qn("w:fldCharType"), "separate")
    t     = OxmlElement("w:t"); t.text = "0"
    fld_e = OxmlElement("w:fldChar"); fld_e.set(qn("w:fldCharType"), "end")
    for el in [fld_b, ins, fld_s, t, fld_e]:
        run._r.append(el)


def add_sidebar_textbox(doc: Document, tipo_peca: str):
    """
    Injeta uma caixa de texto flutuante com o tipo de peça em vertical
    no lado esquerdo da página (no espaço da margem esquerda).
    """
    NS = {
        'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
        'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'mc':  'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    }

    # Dimensões em EMU (1 cm = 360000 EMU)
    box_w  = int(1.2 * 360000)    # largura = 1.2 cm (dimensão "curta" antes de rotar)
    box_h  = int(18.0 * 360000)   # altura = 18 cm (span vertical)
    pos_x  = int(-2.8 * 360000)   # x negativo = dentro da margem esquerda
    pos_y  = int(4.0  * 360000)   # y a partir do topo

    xml = f"""
<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
     xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
     xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
     xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
     xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
  <w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>
  <w:r>
    <w:rPr/>
    <w:drawing>
      <wp:anchor distT="0" distB="0" distL="0" distR="0"
                 simplePos="0" relativeHeight="251659264"
                 behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="margin">
          <wp:posOffset>{pos_x}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="margin">
          <wp:posOffset>{pos_y}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{box_w}" cy="{box_h}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="1" name="SidebarText"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp>
              <wps:cNvSpPr txBx="1"/>
              <wps:spPr>
                <a:xfrm rot="16200000">
                  <a:off x="0" y="0"/>
                  <a:ext cx="{box_h}" cy="{box_w}"/>
                </a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                <a:noFill/>
                <a:ln><a:noFill/></a:ln>
              </wps:spPr>
              <wps:txbx>
                <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                  <w:p>
                    <w:pPr>
                      <w:jc w:val="center"/>
                      <w:spacing w:before="0" w:after="0"/>
                    </w:pPr>
                    <w:r>
                      <w:rPr>
                        <w:color w:val="{NAVY_HEX}"/>
                        <w:sz w:val="18"/>
                        <w:szCs w:val="18"/>
                        <w:spacing w:val="120"/>
                      </w:rPr>
                      <w:t>{tipo_peca}</w:t>
                    </w:r>
                  </w:p>
                </w:txbxContent>
              </wps:txbx>
              <wps:bodyPr rot="0" vert="horz" wrap="square" lIns="0" tIns="0" rIns="0" bIns="0" anchor="ctr"/>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </w:r>
</w:p>"""

    p_elem = etree.fromstring(xml)
    doc.element.body.insert(0, p_elem)


# ─── Template builder ─────────────────────────────────────────────────────────

def criar_template(tipo_peca: str = "AGRAVO DE INSTRUMENTO", info: dict = CONTACT):
    doc = Document()

    # == Margens ==
    section = doc.sections[0]
    section.top_margin    = Cm(4.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(4.0)
    section.right_margin  = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(1.2)

    # == CABEÇALHO ==
    header = section.header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    hp = header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), NAVY2_HEX)
    pBdr.append(bot)
    pPr.append(pBdr)

    # Salva imagem temporária e insere
    header_bytes  = gerar_header_png(info)
    tmp_img_path  = ASSETS / "_header_tmp.png"
    tmp_img_path.write_bytes(header_bytes)
    run = hp.add_run()
    run.add_picture(str(tmp_img_path), width=Cm(16.5))
    tmp_img_path.unlink()

    # == RODAPÉ — barra navy com número de página ==
    footer = section.footer
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    ft = footer.add_table(1, 3, width=Cm(16.5))
    remove_table_borders(ft)
    set_table_full_width(ft)

    # Célula esquerda: AE + nome
    lc = ft.cell(0, 0)
    set_cell_color(lc, NAVY_HEX)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = lp._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "60")
    spacing.set(qn("w:after"),  "60")
    pPr.append(spacing)
    r1 = lp.add_run("AE  ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.name = "Times New Roman"
    r1.font.color.rgb = RGBColor(255, 255, 255)
    r2 = lp.add_run("ANGELO EPIFANIO ADVOCACIA")
    r2.font.size = Pt(7)
    r2.font.color.rgb = RGBColor(180, 200, 220)

    # Célula central: contatos
    cc = ft.cell(0, 1)
    set_cell_color(cc, NAVY_HEX)
    cp = cc.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = cp._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "60")
    spacing.set(qn("w:after"),  "60")
    pPr.append(spacing)
    cr = cp.add_run(f"{info['site']}  |  {info['email']}  |  {info['cidade']}")
    cr.font.size = Pt(6.5)
    cr.font.color.rgb = RGBColor(180, 200, 220)

    # Célula direita: número de página
    rc = ft.cell(0, 2)
    set_cell_color(rc, NAVY_HEX)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = rp._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "60")
    spacing.set(qn("w:after"),  "60")
    pPr.append(spacing)
    rr = rp.add_run()
    rr.font.size = Pt(8)
    rr.font.bold = True
    rr.font.color.rgb = RGBColor(255, 255, 255)
    add_page_field(rr)

    # == Ajuste de largura das células do rodapé ==
    lc.width = Cm(6)
    cc.width = Cm(8)
    rc.width = Cm(2.5)

    # == Parágrafo placeholder no body ==
    doc.add_paragraph()

    # Salva
    out = MODELOS / "AGRAVO_INSTRUMENTO.docx"
    doc.save(str(out))
    print(f"Template base salvo: {out}")

    # Recarrega e adiciona a barra lateral
    doc2 = Document(str(out))
    add_sidebar_textbox(doc2, tipo_peca)
    doc2.save(str(out))
    print(f"Barra lateral adicionada: {out}")

    return out


if __name__ == "__main__":
    criar_template()
