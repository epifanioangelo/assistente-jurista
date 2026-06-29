#!/usr/bin/env python3
"""
Gerador de DOCX para peças jurídicas.
Usa os modelos de "Modelos jus/" como template base, preservando fontes,
estilos (Manuale / Alegreya Sans), margens e numeração de página já configurados.

Uso:
    python3 gerar_docx.py entrada.txt saida.docx [template.docx]

    Se template.docx for omitido, usa "Modelos jus/1._INICIAL.docx" como padrão.

────────────────────────────────────────────────────────────────────────────────
MARCADORES DO ARQUIVO DE ENTRADA
────────────────────────────────────────────────────────────────────────────────
@CENTRO       linha centralizada em negrito  (endereçamento, nome da ação)
@TITULO       capítulo (negrito, centralizado — estilo "1. Título" do template)
@SUBTITULO    subseção primária  (estilo "2. Subtítulo" / "3 Subtítulo")
@SUBTITULO2   subseção secundária
@CITACAO      início de bloco de citação  (estilo "4. Citação")
@FIM          fecha bloco de citação, linha do tempo, tabela ou bullets

@LINHA_DO_TEMPO
Data | Fato | Documento | Observação
12/10/2025 | Diagnóstico confirmado | Laudo — Anexo 1 | Início do quadro
@FIM

@TABELA_QUESTOES          ← 2 colunas: questão | resposta (header automático)
Gratuidade | Renda comprometida
@FIM

@TABELA_PROVAS            ← 3 colunas: fato | documento | anexo (header automático)
Prescrição | Parecer médico | Anexo 1
@FIM

@TABELA_URGENCIA Título   ← 2 colunas: item | referência (header colorido)
Risco de progressão tumoral | Anexo 1
@FIM

@TABELA_LIVRE N_COLUNAS   ← tabela genérica; 1ª linha = cabeçalho
Col1 | Col2 | Col3
Dado | Dado | Dado
@FIM

@BULLETS                  ← lista com marcador
Item um
Item dois
@FIM

@LISTA_ALFA               ← lista a) b) c)  (estilo "5. Lista alfabética")
Primeiro item
Segundo item
@FIM

@TITULO_PECA              ← título da peça (grande, centralizado, com ◆ abaixo)
AGRAVO DE INSTRUMENTO
COM PEDIDO DE TUTELA RECURSAL
@FIM

@SECAO_DESTAQUE Texto     ← seção estilizada com ◆ abaixo (ex: JUSTIFICATIVA DA REFORMA)

@TABELA_SUMARIO Título1 | Título2   ← 2 colunas com cabeçalhos navy
Conteúdo esquerda          ← linhas da coluna esquerda
@COLUNA
Conteúdo direita           ← linhas da coluna direita
@FIM

@IMAGEM caminho/para/imagem.png [largura_cm]   ← insere imagem (padrão: 15cm)
@ESPACO N      ← N linhas em branco (padrão: 1)
@LINHA         ← linha horizontal separadora
@ASSINATURA    ← bloco de assinatura centralizado (após @ASSINATURA)
@AVISO         ← IGNORADO — avisos de IA não constam nas peças processuais
────────────────────────────────────────────────────────────────────────────────
"""

import sys
import re
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY_RGB  = RGBColor(13, 35, 64)
NAVY2_RGB = RGBColor(26, 58, 92)


# ── Cores das tabelas (extraídas dos modelos) ─────────────────────────────────
COR_HEADER    = 'D9D9D9'   # cinza médio — cabeçalho de tabela
COR_LINHA_A   = 'F8F8F6'   # creme muito claro — linhas pares
COR_LINHA_B   = 'FFFFFF'   # branco — linhas ímpares
COR_DESTAQUE  = 'F1F1EB'   # bege — células de destaque / sub-header


# ── Mapa de estilos do template ───────────────────────────────────────────────
# Nomes exatos dos estilos presentes nos modelos "Modelos jus/"
STYLE_TITULO     = '1. Título'
STYLE_SUBTITULO  = '3 Subtítulo'
STYLE_SUBTITULO2 = '3 Subtítulo secundário'
STYLE_CORPO      = '3 Corpo do texto'
STYLE_CITACAO    = '4. Citação'
STYLE_LISTA_ALFA = '5. Lista alfabética'
STYLE_BULLETS    = '6 Bullets'


# ── Helpers XML ───────────────────────────────────────────────────────────────
def set_cell_color(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_bold(cell, bold: bool = True):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = bold


def add_table_border(table):
    """Borda simples em todas as células."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'AAAAAA')
        borders.append(el)
    tblPr.append(borders)


def table_full_width(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)


# ── Inline markdown (bold / italic) ──────────────────────────────────────────
def parse_inline(texto: str):
    partes = []
    pattern = re.compile(r'\*\*(.*?)\*\*|\*(.*?)\*')
    pos = 0
    for m in pattern.finditer(texto):
        if m.start() > pos:
            partes.append({'t': texto[pos:m.start()], 'b': False, 'i': False})
        if m.group(1) is not None:
            partes.append({'t': m.group(1), 'b': True,  'i': False})
        else:
            partes.append({'t': m.group(2), 'b': False, 'i': True})
        pos = m.end()
    if pos < len(texto):
        partes.append({'t': texto[pos:], 'b': False, 'i': False})
    return partes


def fill_para(para, texto: str, bold: bool = False):
    for seg in parse_inline(texto):
        run = para.add_run(seg['t'])
        run.font.bold   = bold or seg['b']
        run.font.italic = seg['i']


def fill_cell(cell, texto: str, bold: bool = False, center: bool = False):
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fill_para(para, texto.strip(), bold=bold)


# ── Abertura do template ──────────────────────────────────────────────────────
def abrir_template(template_path: Path) -> Document:
    """Abre o template e limpa o conteúdo preservando estilos e seções."""
    doc = Document(str(template_path))
    body = doc.element.body
    # Guarda o último sectPr (controla margens, cabeçalho, rodapé)
    sect_pr = body.find(qn('w:sectPr'))
    # Remove tudo menos sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)
    # Garante um parágrafo inicial vazio (python-docx exige)
    p = OxmlElement('w:p')
    body.insert(0, p)
    return doc


def style_exists(doc: Document, name: str) -> bool:
    return any(s.name == name for s in doc.styles)


def safe_style(doc: Document, preferred: str, fallback: str = 'Normal') -> str:
    return preferred if style_exists(doc, preferred) else fallback


# ── Tabelas ───────────────────────────────────────────────────────────────────
def build_linha_do_tempo(doc: Document, linhas: list[str]):
    """4 colunas: Data | Fato | Documento | Observação"""
    headers = ['Data', 'Fato', 'Documento', 'Observação']
    table = doc.add_table(rows=1, cols=4)
    table_full_width(table)
    add_table_border(table)
    # Cabeçalho
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_color(cell, COR_HEADER)
        fill_cell(cell, h, bold=True, center=True)
    # Dados
    for ri, linha in enumerate(linhas):
        cols = [c.strip() for c in linha.split('|')]
        while len(cols) < 4:
            cols.append('')
        row = table.add_row()
        cor = COR_LINHA_A if ri % 2 == 0 else COR_LINHA_B
        for ci, val in enumerate(cols[:4]):
            set_cell_color(row.cells[ci], cor)
            fill_cell(row.cells[ci], val)
    doc.add_paragraph()


def build_tabela_questoes(doc: Document, linhas: list[str]):
    """2 colunas: Questão processual | Resposta"""
    table = doc.add_table(rows=1, cols=2)
    table_full_width(table)
    add_table_border(table)
    hrow = table.rows[0]
    set_cell_color(hrow.cells[0], COR_HEADER)
    set_cell_color(hrow.cells[1], COR_HEADER)
    fill_cell(hrow.cells[0], 'Questão processual', bold=True, center=True)
    fill_cell(hrow.cells[1], 'Posição da parte', bold=True, center=True)
    for ri, linha in enumerate(linhas):
        cols = [c.strip() for c in linha.split('|')]
        while len(cols) < 2:
            cols.append('')
        row = table.add_row()
        cor = COR_LINHA_A if ri % 2 == 0 else COR_LINHA_B
        set_cell_color(row.cells[0], cor)
        set_cell_color(row.cells[1], COR_LINHA_B)
        fill_cell(row.cells[0], cols[0], bold=True)
        fill_cell(row.cells[1], cols[1])
    doc.add_paragraph()


def build_tabela_provas(doc: Document, linhas: list[str]):
    """3 colunas: Fato alegado | Documento | Anexo"""
    table = doc.add_table(rows=1, cols=3)
    table_full_width(table)
    add_table_border(table)
    hrow = table.rows[0]
    for ci, h in enumerate(['Fato alegado', 'Documento', 'Anexo']):
        set_cell_color(hrow.cells[ci], COR_HEADER)
        fill_cell(hrow.cells[ci], h, bold=True, center=True)
    for ri, linha in enumerate(linhas):
        cols = [c.strip() for c in linha.split('|')]
        while len(cols) < 3:
            cols.append('')
        row = table.add_row()
        cor = COR_LINHA_A if ri % 2 == 0 else COR_LINHA_B
        for ci in range(3):
            set_cell_color(row.cells[ci], cor)
            fill_cell(row.cells[ci], cols[ci] if ci < len(cols) else '')
    doc.add_paragraph()


def build_tabela_urgencia(doc: Document, titulo: str, linhas: list[str]):
    """2 colunas com sub-header colorido: item | referência"""
    table = doc.add_table(rows=1, cols=2)
    table_full_width(table)
    add_table_border(table)
    # Sub-header (merge)
    hrow = table.rows[0]
    hrow.cells[0].merge(hrow.cells[1])
    set_cell_color(hrow.cells[0], COR_DESTAQUE)
    fill_cell(hrow.cells[0], titulo, bold=True, center=True)
    for ri, linha in enumerate(linhas):
        cols = [c.strip() for c in linha.split('|')]
        while len(cols) < 2:
            cols.append('')
        row = table.add_row()
        cor = COR_LINHA_A if ri % 2 == 0 else COR_LINHA_B
        set_cell_color(row.cells[0], cor)
        set_cell_color(row.cells[1], cor)
        fill_cell(row.cells[0], cols[0])
        fill_cell(row.cells[1], cols[1] if len(cols) > 1 else '')
    doc.add_paragraph()


def set_run_spacing(run, half_pts: int):
    """Define espaçamento entre caracteres (letter-spacing) em meio-pontos."""
    rPr = run._r.get_or_add_rPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:val'), str(half_pts))
    rPr.append(sp)


def build_titulo_peca(doc: Document, linhas: list[str]):
    """Título da peça processual — grande, centralizado, negrito, com ◆ abaixo."""
    for linha in linhas:
        if not linha.strip():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(linha.strip().upper())
        r.font.bold  = True
        r.font.size  = Pt(16)
        r.font.name  = 'Times New Roman'
        r.font.color.rgb = NAVY_RGB
    # ◆ abaixo do título
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('◆')
    r2.font.size = Pt(10)
    r2.font.color.rgb = NAVY_RGB
    doc.add_paragraph()


def add_secao_destaque(doc: Document, texto: str):
    """Seção de destaque — ex: JUSTIFICATIVA DA REFORMA — com ◆ abaixo."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto.strip().upper())
    r.font.bold  = False
    r.font.size  = Pt(11)
    r.font.name  = 'Times New Roman'
    r.font.color.rgb = NAVY_RGB
    set_run_spacing(r, 80)   # ~4 pt de letter-spacing
    # ◆
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('◆')
    r2.font.size = Pt(8)
    r2.font.color.rgb = NAVY_RGB
    doc.add_paragraph()


def _header_cell_navy(cell, texto: str):
    """Preenche célula com header estilizado: texto tracked + ◆ abaixo."""
    set_cell_color(cell, '0D2340')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '60')
    sp.set(qn('w:after'),  '40')
    pPr.append(sp)
    r = p.add_run(texto.strip().upper())
    r.font.bold  = False
    r.font.size  = Pt(9)
    r.font.color.rgb = RGBColor(255, 255, 255)
    set_run_spacing(r, 80)
    # ◆ em parágrafo próprio dentro da célula
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr2 = p2._p.get_or_add_pPr()
    sp2  = OxmlElement('w:spacing')
    sp2.set(qn('w:before'), '0')
    sp2.set(qn('w:after'),  '60')
    pPr2.append(sp2)
    r2 = p2.add_run('◆')
    r2.font.size = Pt(7)
    r2.font.color.rgb = RGBColor(200, 210, 225)


def build_tabela_sumario(doc: Document, titulo_esq: str, titulo_dir: str,
                          linhas_esq: list[str], linhas_dir: list[str]):
    """
    2 colunas com cabeçalhos navy e losango.
    Usado para DECISÃO RECORRIDA | PEDIDO e similares.
    """
    table = doc.add_table(rows=1, cols=2)
    table_full_width(table)
    add_table_border(table)

    hrow = table.rows[0]
    _header_cell_navy(hrow.cells[0], titulo_esq)
    _header_cell_navy(hrow.cells[1], titulo_dir)

    # Junta pares de linhas (esq e dir podem ter comprimentos diferentes)
    max_rows = max(len(linhas_esq), len(linhas_dir))
    for ri in range(max_rows):
        row = table.add_row()
        cor = 'F8F8F6' if ri % 2 == 0 else 'FFFFFF'
        set_cell_color(row.cells[0], cor)
        set_cell_color(row.cells[1], cor)
        esq = linhas_esq[ri].strip() if ri < len(linhas_esq) else ''
        dir_ = linhas_dir[ri].strip() if ri < len(linhas_dir) else ''
        fill_cell(row.cells[0], esq)
        fill_cell(row.cells[1], dir_)
    doc.add_paragraph()


def build_tabela_livre(doc: Document, n_cols: int, linhas: list[str]):
    """Tabela genérica: primeira linha = cabeçalho."""
    if not linhas:
        return
    table = doc.add_table(rows=0, cols=n_cols)
    table_full_width(table)
    add_table_border(table)
    for ri, linha in enumerate(linhas):
        cols = [c.strip() for c in linha.split('|')]
        while len(cols) < n_cols:
            cols.append('')
        row = table.add_row()
        is_header = ri == 0
        cor = COR_HEADER if is_header else (COR_LINHA_A if ri % 2 == 0 else COR_LINHA_B)
        for ci in range(n_cols):
            set_cell_color(row.cells[ci], cor)
            fill_cell(row.cells[ci], cols[ci] if ci < len(cols) else '',
                      bold=is_header, center=is_header)
    doc.add_paragraph()


# ── Parser principal ──────────────────────────────────────────────────────────
def parse(doc: Document, texto: str):
    linhas = texto.splitlines()
    i = 0
    modo           = None
    buf_tabela     = []
    tab_titulo     = ''
    tab_ncols      = 2
    tab_col_esq    = ''
    tab_col_dir    = ''
    buf_sumario_esq = []
    buf_sumario_dir = []
    _sumario_lado   = 'esq'

    while i < len(linhas):
        linha = linhas[i].rstrip()
        stripped = linha.strip()

        # ── Fechamento de bloco ──────────────────────────────────────────────
        if stripped == '@FIM':
            if modo == 'linha_do_tempo':
                build_linha_do_tempo(doc, buf_tabela)
            elif modo == 'tabela_questoes':
                build_tabela_questoes(doc, buf_tabela)
            elif modo == 'tabela_provas':
                build_tabela_provas(doc, buf_tabela)
            elif modo == 'tabela_urgencia':
                build_tabela_urgencia(doc, tab_titulo, buf_tabela)
            elif modo == 'tabela_livre':
                build_tabela_livre(doc, tab_ncols, buf_tabela)
            elif modo == 'titulo_peca':
                build_titulo_peca(doc, buf_tabela)
            elif modo == 'tabela_sumario':
                build_tabela_sumario(doc, tab_col_esq, tab_col_dir,
                                      buf_sumario_esq, buf_sumario_dir)
            elif modo in ('bullets', 'lista_alfa'):
                pass  # já adicionado linha a linha
            modo = None
            buf_tabela = []
            i += 1
            continue

        # ── Separador de colunas da @TABELA_SUMARIO ─────────────────────────
        if modo == 'tabela_sumario' and stripped == '@COLUNA':
            _sumario_lado = 'dir'
            i += 1
            continue

        if modo == 'tabela_sumario':
            if stripped:
                if _sumario_lado == 'esq':
                    buf_sumario_esq.append(stripped)
                else:
                    buf_sumario_dir.append(stripped)
            i += 1
            continue

        if modo == 'titulo_peca':
            buf_tabela.append(linha)
            i += 1
            continue

        # ── Dentro de bloco de tabela/lista ─────────────────────────────────
        if modo in ('linha_do_tempo', 'tabela_questoes', 'tabela_provas',
                    'tabela_urgencia', 'tabela_livre'):
            if stripped:
                buf_tabela.append(stripped)
            i += 1
            continue

        if modo == 'citacao':
            if stripped:
                p = doc.add_paragraph(style=safe_style(doc, STYLE_CITACAO))
                fill_para(p, stripped)
            i += 1
            continue

        if modo == 'aviso':
            i += 1
            continue

        if modo == 'bullets':
            if stripped:
                p = doc.add_paragraph(style=safe_style(doc, STYLE_BULLETS))
                fill_para(p, stripped)
            i += 1
            continue

        if modo == 'lista_alfa':
            if stripped:
                p = doc.add_paragraph(style=safe_style(doc, STYLE_LISTA_ALFA))
                fill_para(p, stripped)
            i += 1
            continue

        if modo == 'assinatura':
            if stripped == '':
                doc.add_paragraph()
            else:
                bold = not re.match(r'^OAB/', stripped)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fill_para(p, stripped, bold=bold)
            i += 1
            continue

        # ── Abertura de blocos ───────────────────────────────────────────────
        if stripped == '@CITACAO':
            modo = 'citacao';  i += 1;  continue
        if stripped == '@BULLETS':
            modo = 'bullets';  i += 1;  continue
        if stripped == '@LISTA_ALFA':
            modo = 'lista_alfa';  i += 1;  continue
        if stripped == '@ASSINATURA':
            modo = 'assinatura'
            doc.add_paragraph()
            doc.add_paragraph()
            i += 1;  continue
        if stripped == '@AVISO':
            modo = 'aviso';  i += 1;  continue
        if stripped == '@LINHA_DO_TEMPO':
            modo = 'linha_do_tempo';  buf_tabela = [];  i += 1;  continue
        if stripped == '@TABELA_QUESTOES':
            modo = 'tabela_questoes';  buf_tabela = [];  i += 1;  continue
        if stripped == '@TABELA_PROVAS':
            modo = 'tabela_provas';  buf_tabela = [];  i += 1;  continue
        m = re.match(r'^@TABELA_URGENCIA\s*(.*)', stripped)
        if m:
            modo = 'tabela_urgencia'
            tab_titulo = m.group(1).strip()
            buf_tabela = [];  i += 1;  continue
        m = re.match(r'^@TABELA_LIVRE\s*(\d+)', stripped)
        if m:
            modo = 'tabela_livre'
            tab_ncols = int(m.group(1))
            buf_tabela = [];  i += 1;  continue

        if stripped == '@TITULO_PECA':
            modo = 'titulo_peca';  buf_tabela = [];  i += 1;  continue

        m = re.match(r'^@TABELA_SUMARIO\s*(.+)', stripped)
        if m:
            partes = m.group(1).split('|')
            tab_col_esq     = partes[0].strip() if len(partes) > 0 else ''
            tab_col_dir     = partes[1].strip() if len(partes) > 1 else ''
            buf_sumario_esq = []
            buf_sumario_dir = []
            _sumario_lado   = 'esq'
            modo = 'tabela_sumario';  i += 1;  continue

        m = re.match(r'^@SECAO_DESTAQUE\s*(.*)', stripped)
        if m:
            add_secao_destaque(doc, m.group(1))
            i += 1;  continue

        # ── Diretivas de parágrafo único ─────────────────────────────────────
        if stripped == '':
            doc.add_paragraph()
            i += 1;  continue

        if re.match(r'^-{3,}$', stripped):
            i += 1;  continue

        if stripped == '@LINHA':
            p = doc.add_paragraph()
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'AAAAAA')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1;  continue

        m = re.match(r'^@ESPACO\s*(\d*)', stripped)
        if m:
            n = int(m.group(1)) if m.group(1) else 1
            for _ in range(n):
                doc.add_paragraph()
            i += 1;  continue

        m = re.match(r'^@IMAGEM\s+(\S+)(?:\s+([\d.]+))?', stripped)
        if m:
            img_path = Path(m.group(1))
            largura  = float(m.group(2)) if m.group(2) else 15.0
            if img_path.exists():
                doc.add_picture(str(img_path), width=Cm(largura))
            else:
                p = doc.add_paragraph(f'[IMAGEM NÃO ENCONTRADA: {img_path}]')
            i += 1;  continue

        if stripped.startswith('@CENTRO'):
            conteudo = stripped[7:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fill_para(p, conteudo, bold=True)
            i += 1;  continue

        if stripped.startswith('@TITULO'):
            conteudo = stripped[7:].strip().upper()
            p = doc.add_paragraph(style=safe_style(doc, STYLE_TITULO))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fill_para(p, conteudo, bold=True)
            i += 1;  continue

        if stripped.startswith('@SUBTITULO2'):
            conteudo = stripped[11:].strip()
            p = doc.add_paragraph(style=safe_style(doc, STYLE_SUBTITULO2, STYLE_SUBTITULO))
            fill_para(p, conteudo)
            i += 1;  continue

        if stripped.startswith('@SUBTITULO'):
            conteudo = stripped[10:].strip()
            p = doc.add_paragraph(style=safe_style(doc, STYLE_SUBTITULO))
            fill_para(p, conteudo, bold=True)
            i += 1;  continue

        # Item de lista avulso a) b) c)
        if re.match(r'^[a-z]\)', stripped):
            p = doc.add_paragraph(style=safe_style(doc, STYLE_LISTA_ALFA))
            fill_para(p, stripped)
            i += 1;  continue

        # Parágrafo normal
        p = doc.add_paragraph(style=safe_style(doc, STYLE_CORPO))
        fill_para(p, stripped)
        i += 1


# ── Entry point ───────────────────────────────────────────────────────────────
def gerar(entrada: Path, saida: Path, template: Path):
    doc = abrir_template(template)
    texto = entrada.read_text(encoding='utf-8')
    parse(doc, texto)
    doc.save(str(saida))
    print(f'DOCX gerado: {saida}')


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    entrada  = Path(sys.argv[1])
    saida    = Path(sys.argv[2])
    base     = Path(__file__).parent.parent   # raiz do projeto
    template = Path(sys.argv[3]) if len(sys.argv) > 3 else base / 'Modelos jus' / '1._INICIAL.docx'

    if not template.exists():
        print(f'Template não encontrado: {template}')
        sys.exit(1)

    gerar(entrada, saida, template)
