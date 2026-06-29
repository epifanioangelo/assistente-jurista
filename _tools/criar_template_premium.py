#!/usr/bin/env python3
"""
Template premium Angelo Epifanio Advocacia.
Construído inteiramente com recursos nativos do Word — sem imagens.
Saída: Modelos jus/AGRAVO_INSTRUMENTO.docx

Uso: python3 criar_template_premium.py
"""

from pathlib import Path
from lxml import etree

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT    = Path(__file__).parent.parent
MODELOS = ROOT / "Modelos jus"

# ── Paleta de cores ───────────────────────────────────────────────────────────
NAVY1 = "0B2341"   # azul primário
NAVY2 = "173A63"   # azul secundário
GRAY  = "8A9099"   # cinza
GOLD  = "B69A54"   # dourado
WHITE = "FFFFFF"

C_NAVY1 = RGBColor(11,  35,  65)
C_NAVY2 = RGBColor(23,  58,  99)
C_GRAY  = RGBColor(138, 144, 153)
C_GOLD  = RGBColor(182, 154, 84)
C_WHITE = RGBColor(255, 255, 255)

# ── Contato (editar conforme necessário) ──────────────────────────────────────
OAB    = "OAB/SP XXXXX"
SITE   = "www.angeloepifanio.com.br"
EMAIL  = "contato@angeloepifanio.com.br"
TEL    = "(12) 99999-9999"
CIDADE = "São José dos Campos/SP"


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS XML
# ═══════════════════════════════════════════════════════════════════════════════

def _sub(parent, tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), str(v))
    parent.append(el)
    return el


def set_cell_color(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_no_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _sub(bdr, f"w:{side}", **{"w:val": "none", "w:sz": "0", "w:color": "auto"})
    tcPr.append(bdr)


def set_cell_border(cell, sides: dict):
    """sides = {'left': ('single', '4', GOLD), ...}"""
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement("w:tcBorders")
    for side, (val, sz, color) in sides.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   val)
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:color"), color)
        bdr.append(el)
    tcPr.append(bdr)


def set_table_no_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    bdr = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        _sub(bdr, f"w:{side}", **{"w:val": "none", "w:sz": "0", "w:color": "auto"})
    tblPr.append(bdr)


def set_table_width(table, width_cm: float):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    # Remove existing width
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    # width in twips (1 cm = 567 twips)
    w_twips = int(width_cm * 567)
    _sub(tblPr, "w:tblW", **{"w:w": str(w_twips), "w:type": "dxa"})


def set_para_spacing(para, before_pt=0, after_pt=6, line_rule="auto", line=276):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement("w:spacing")
    sp.set(qn("w:before"),    str(int(before_pt * 20)))
    sp.set(qn("w:after"),     str(int(after_pt  * 20)))
    sp.set(qn("w:lineRule"),  line_rule)
    sp.set(qn("w:line"),      str(line))   # 276 = 1.15×; 360 = 1.5×
    pPr.append(sp)


def set_para_border_bottom(para, color: str, sz_pts: float = 0.5, space: int = 4):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(max(1, int(sz_pts * 8))))   # 8ths of a pt
    bot.set(qn("w:space"), str(space))
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_run_spacing(run, twips: int):
    """Letter-spacing: 20 twips = 1pt."""
    rPr = run._r.get_or_add_rPr()
    sp  = OxmlElement("w:spacing")
    sp.set(qn("w:val"), str(twips))
    rPr.append(sp)


def add_run_small_caps(run):
    rPr = run._r.get_or_add_rPr()
    OxmlElement("w:smallCaps")
    sc = OxmlElement("w:smallCaps")
    rPr.append(sc)


def add_page_field(run):
    for tag, text in [("begin", None), ("instrText", " PAGE "), ("separate", None),
                      ("t", "0"), ("end", None)]:
        if tag in ("instrText", "t"):
            el = OxmlElement(f"w:{tag}")
            if tag == "instrText":
                el.set(qn("xml:space"), "preserve")
            el.text = text
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        run._r.append(el)


def set_cell_vert_align(cell, align="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    va   = OxmlElement("w:vAlign")
    va.set(qn("w:val"), align)
    tcPr.append(va)


def set_table_indent(table, indent_cm: float):
    """Recua tabela (negativo = empurra para a esquerda, além da margem)."""
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"),    str(int(indent_cm * 567)))
    ind.set(qn("w:type"), "dxa")
    tblPr.append(ind)


# ═══════════════════════════════════════════════════════════════════════════════
# ESTILOS
# ═══════════════════════════════════════════════════════════════════════════════

def criar_estilos(doc: Document):
    """Cria os estilos AE_ usados pelo gerar_docx.py e pelo template."""

    def _make(name, base=None, font_name="Georgia", size_pt=12,
              bold=False, italic=False, color=None, align=None,
              space_before=0, space_after=6, line=360):
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        if base and base in [s.name for s in doc.styles]:
            st.base_style = doc.styles[base]
        f = st.font
        f.name  = font_name
        f.size  = Pt(size_pt)
        f.bold  = bold
        f.italic = italic
        if color:
            f.color.rgb = color
        pf = st.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after  = Pt(space_after)
        pf.line_spacing = Pt(line / 20)  # linha em pontos
        if align:
            pf.alignment = align
        return st

    _make("AE Título 1", font_name="Georgia", size_pt=14, bold=True,
          color=C_NAVY1, align=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=12, space_after=6, line=360)

    _make("AE Título 2", font_name="Georgia", size_pt=12, bold=True,
          color=C_NAVY1, space_before=10, space_after=4, line=360)

    _make("AE Título 3", font_name="Georgia", size_pt=11, bold=True, italic=True,
          color=C_NAVY2, space_before=8, space_after=4, line=360)

    _make("AE Corpo", font_name="Times New Roman", size_pt=12,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6, line=360)

    _make("AE Citação", font_name="Times New Roman", size_pt=11, italic=True,
          color=C_NAVY2, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_before=6, space_after=6, line=340)

    _make("AE Nota", font_name="Arial", size_pt=9,
          color=C_GRAY, space_before=0, space_after=3, line=240)

    _make("AE Endereçamento", font_name="Times New Roman", size_pt=12, bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, line=360)

    _make("AE Ornamento", font_name="Georgia", size_pt=10,
          color=C_GOLD, align=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=4, space_after=4, line=240)


# ═══════════════════════════════════════════════════════════════════════════════
# CABEÇALHO
# ═══════════════════════════════════════════════════════════════════════════════

def criar_cabecalho(section):
    """
    Cabeçalho nativo Word:
      Coluna esquerda (vazia) | Centro (logo AE) | Direita (contato)
    """
    header = section.header
    # Remove parágrafo padrão
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    # Área de texto: 16.5 cm = 9355 twips
    # Células: 1.5cm(851) + 10.0cm(5669) + 5.0cm(2835) = 9355
    tab = header.add_table(1, 3, width=Cm(16.5))
    set_table_no_borders(tab)
    _set_table_xml(tab, 9355, 0, [851, 5669, 2835])

    lc = tab.cell(0, 0)   # Esquerda vazia
    cc = tab.cell(0, 1)   # Centro — logo
    rc = tab.cell(0, 2)   # Direita — contato

    for cell in (lc, cc, rc):
        set_cell_no_borders(cell)
        set_cell_vert_align(cell, "center")

    # ── Logo central ──────────────────────────────────────────────────────────
    # "AE" em Georgia Bold grande
    p_ae = cc.paragraphs[0]
    p_ae.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p_ae, before_pt=4, after_pt=2, line_rule="auto", line=240)
    r_ae = p_ae.add_run("AE")
    r_ae.font.name  = "Georgia"
    r_ae.font.bold  = True
    r_ae.font.size  = Pt(28)
    r_ae.font.color.rgb = C_NAVY1
    add_run_spacing(r_ae, 80)   # ~4pt letter-spacing

    # Linha dourada sob o AE
    set_para_border_bottom(p_ae, GOLD, sz_pts=0.5, space=3)

    # "ANGELO EPIFANIO"
    p_nm = cc.add_paragraph()
    p_nm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p_nm, before_pt=4, after_pt=0, line_rule="auto", line=240)
    r_nm = p_nm.add_run("ANGELO EPIFANIO")
    r_nm.font.name  = "Georgia"
    r_nm.font.bold  = True
    r_nm.font.size  = Pt(9)
    r_nm.font.color.rgb = C_NAVY1
    add_run_spacing(r_nm, 100)

    # Linha dourada
    set_para_border_bottom(p_nm, GOLD, sz_pts=0.5, space=3)

    # "ADVOCACIA"
    p_adv = cc.add_paragraph()
    p_adv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p_adv, before_pt=3, after_pt=4, line_rule="auto", line=240)
    r_adv = p_adv.add_run("ADVOCACIA")
    r_adv.font.name  = "Arial"
    r_adv.font.size  = Pt(7.5)
    r_adv.font.color.rgb = C_GRAY
    add_run_spacing(r_adv, 160)

    # ── Contato direito ───────────────────────────────────────────────────────
    contatos = [
        (OAB,   9.0, C_NAVY1, True,  80),
        ("",    0.0, C_GRAY,  False, 0),    # espaço
        (SITE,  7.5, C_GRAY,  False, 0),
        (EMAIL, 7.5, C_GRAY,  False, 0),
        (TEL,   7.5, C_GRAY,  False, 0),
    ]
    first = True
    for txt, size, color, bold, spacing in contatos:
        if first:
            p = rc.paragraphs[0]; first = False
        else:
            p = rc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_para_spacing(p, before_pt=0, after_pt=1, line_rule="auto", line=240)
        if txt:
            r = p.add_run(txt)
            r.font.name  = "Arial"
            r.font.size  = Pt(size)
            r.font.bold  = bold
            r.font.color.rgb = color
            if spacing:
                add_run_spacing(r, spacing)


# ═══════════════════════════════════════════════════════════════════════════════
# RODAPÉ
# ═══════════════════════════════════════════════════════════════════════════════

def _set_table_xml(table, page_w_twips: int, indent_twips: int,
                   cell_widths_twips: list[int]):
    """
    Define largura, recuo e largura de células em twips via XML direto.
    Evita conflitos de unidades entre python-docx e Word.
    """
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)

    # Remove tags conflitantes
    for tag in (qn("w:tblW"), qn("w:tblInd"), qn("w:tblLayout"), qn("w:jc")):
        for old in list(tblPr.findall(tag)):
            tblPr.remove(old)

    # Largura total absoluta
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(page_w_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Recuo (negativo = entra na margem esquerda)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"),    str(indent_twips))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

    # Layout fixo (Word respeita as larguras definidas)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # Largura de cada célula
    for cell, w in zip(table.row_cells(0), cell_widths_twips):
        tcPr = cell._tc.get_or_add_tcPr()
        for old in list(tcPr.findall(qn("w:tcW"))):
            tcPr.remove(old)
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"),    str(w))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)


def _add_tab_stop(pPr, pos_twips: int, align: str):
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    t = OxmlElement("w:tab")
    t.set(qn("w:val"), align)
    t.set(qn("w:pos"), str(pos_twips))
    tabs.append(t)


def _set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0):
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)


def _run(p, txt, size, color, bold=False, spacing=0):
    r = p.add_run(txt)
    r.font.name      = "Arial"
    r.font.size      = Pt(size)
    r.font.color.rgb = color
    r.font.bold      = bold
    if spacing:
        add_run_spacing(r, spacing)
    return r


def criar_rodape(section):
    """
    Rodape: paragrafo com indent negativo (esq + dir) para cobertura total da pagina.
    Altura = 2 linhas × 567 twips/linha = 1134 twips = 2cm com lineRule=exact.
    Layout: AE (esq) | nome+site (centro) | email+cidade (dir) via tab stops.
    """
    footer = section.footer
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    # Medidas em twips (1cm = 567 twips)
    LEFT_MG  = 1417   # margem esquerda 2.5cm
    RIGHT_MG = 1134   # margem direita  2.0cm
    TEXT_W   = 9355   # area de texto   16.5cm
    PAGE_W   = TEXT_W + LEFT_MG + RIGHT_MG  # 11906 twips = 21cm

    # ── Paragrafo principal ───────────────────────────────────────────────────
    fp  = footer.add_paragraph()
    pPr = fp._p.get_or_add_pPr()

    # Shading navy
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), NAVY1)
    pPr.append(shd)

    # Indent negativo em ambos os lados → cobre toda a largura da pagina
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"),  str(-LEFT_MG))
    ind.set(qn("w:right"), str(-RIGHT_MG))
    pPr.append(ind)

    # 2 linhas × 567 twips = 1134 twips (2cm) de altura total
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "0")
    sp.set(qn("w:line"), "567"); sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)

    # Tab stops medidos a partir da borda esquerda da area de texto (pos=0)
    # Centro da pagina:  -LEFT_MG + PAGE_W/2 = -1417 + 5953 = 4536
    # Direita c/ padding: PAGE_W - LEFT_MG - 270 = 11906 - 1417 - 270 = 10219
    tabs = OxmlElement("w:tabs")
    for pos, align in [(4536, "center"), (10219, "right")]:
        t = OxmlElement("w:tab")
        t.set(qn("w:val"), align); t.set(qn("w:pos"), str(pos))
        tabs.append(t)
    pPr.append(tabs)

    # Helper local
    def ar(txt, size, color, bold=False, spacing=0):
        r = fp.add_run(txt)
        r.font.name = "Arial"; r.font.size = Pt(size)
        r.font.color.rgb = color; r.font.bold = bold
        if spacing: add_run_spacing(r, spacing)
        return r

    # ── LINHA 1: AE  |  ANGELO EPIFANIO ADVOCACIA  |  contato@... ────────────
    ar("   AE", 16, C_WHITE, bold=True, spacing=80)
    ar("\t", 7, C_WHITE)
    ar("ANGELO EPIFANIO ADVOCACIA", 8, C_WHITE, bold=True, spacing=15)
    ar("\t", 7, C_WHITE)
    ar(EMAIL + "  ", 7, RGBColor(160, 185, 210))

    # Quebra de linha suave (soft return dentro do mesmo paragrafo)
    br_r = fp.add_run()
    br_r._r.append(OxmlElement("w:br"))

    # ── LINHA 2: (vazio)  |  www.site.com  |  São José dos Campos/SP ─────────
    ar("", 7, C_WHITE)
    ar("\t", 7, C_WHITE)
    ar(SITE, 7, RGBColor(160, 185, 210))
    ar("\t", 7, C_WHITE)
    ar(CIDADE + "  ", 7, RGBColor(160, 185, 210))

# ═══════════════════════════════════════════════════════════════════════════════
# BARRA LATERAL (linha dourada + texto vertical)
# ═══════════════════════════════════════════════════════════════════════════════

def adicionar_sidebar(doc: Document, tipo_peca: str):
    """
    Injeta duas formas flutuantes via XML:
      1. Linha dourada fina (retângulo preenchido, 0.5pt de largura)
      2. Caixa de texto com texto vertical em azul claro
    Ambas ancoradas na margem esquerda.
    """
    # EMU: 1 cm = 360000
    cm = 360000

    # Dimensões
    linha_w  = 6350          # 0.5 pt de largura
    linha_h  = int(20 * cm)  # 20 cm de altura
    box_w    = int(1.0 * cm)  # 1 cm de largura
    box_h    = linha_h

    # Posições (relativas à margem)
    pos_x_linha = int(-1.8 * cm)   # na margem esquerda
    pos_x_text  = int(-1.7 * cm)   # ao lado da linha (ligeiramente à dir)
    pos_y       = int(3.5 * cm)    # abaixo do cabeçalho

    WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
    WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def anchor_xml(doc_id, name, cx, cy, pos_x, pos_y, content_xml, behind=False):
        return f"""
<w:p xmlns:w="{W}" xmlns:wp="{WP}" xmlns:a="{A}" xmlns:wps="{WPS}">
  <w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>
  <w:r>
    <w:drawing>
      <wp:anchor distT="0" distB="0" distL="0" distR="0"
                 simplePos="0" relativeHeight="251659264"
                 behindDoc="{'1' if behind else '0'}" locked="0"
                 layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="margin">
          <wp:posOffset>{pos_x}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="topMargin">
          <wp:posOffset>{pos_y}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="{doc_id}" name="{name}"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
          <a:graphicData uri="{WPS.replace('http','http')}">
            {content_xml}
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </w:r>
</w:p>"""

    # 1. Linha dourada (retângulo preenchido, sem borda)
    gold_shape = f"""
<wps:wsp xmlns:wps="{WPS}">
  <wps:cNvSpPr txBx="0"/>
  <wps:spPr>
    <a:xfrm xmlns:a="{A}">
      <a:off x="0" y="0"/>
      <a:ext cx="{linha_w}" cy="{linha_h}"/>
    </a:xfrm>
    <a:prstGeom xmlns:a="{A}" prst="rect"><a:avLst/></a:prstGeom>
    <a:solidFill xmlns:a="{A}"><a:srgbClr val="{GOLD}"/></a:solidFill>
    <a:ln xmlns:a="{A}"><a:noFill/></a:ln>
  </wps:spPr>
  <wps:bodyPr/>
</wps:wsp>"""

    # 2. Caixa de texto com texto vertical
    text_shape = f"""
<wps:wsp xmlns:wps="{WPS}">
  <wps:cNvSpPr txBx="1"/>
  <wps:spPr>
    <a:xfrm xmlns:a="{A}">
      <a:off x="0" y="0"/>
      <a:ext cx="{box_w}" cy="{box_h}"/>
    </a:xfrm>
    <a:prstGeom xmlns:a="{A}" prst="rect"><a:avLst/></a:prstGeom>
    <a:noFill xmlns:a="{A}"/>
    <a:ln xmlns:a="{A}"><a:noFill/></a:ln>
  </wps:spPr>
  <wps:txbx>
    <w:txbxContent xmlns:w="{W}">
      <w:p>
        <w:pPr>
          <w:jc w:val="center"/>
          <w:spacing w:before="0" w:after="0"/>
        </w:pPr>
        <w:r>
          <w:rPr>
            <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
            <w:color w:val="{NAVY2}"/>
            <w:sz w:val="14"/>
            <w:szCs w:val="14"/>
            <w:spacing w:val="120"/>
          </w:rPr>
          <w:t>{tipo_peca}</w:t>
        </w:r>
      </w:p>
    </w:txbxContent>
  </wps:txbx>
  <wps:bodyPr rot="0" vert="vert270" wrap="square"
              lIns="72000" tIns="72000" rIns="72000" bIns="72000"
              anchor="ctr" anchorCtr="0"/>
</wps:wsp>"""

    xml_linha = anchor_xml(10, "GoldLine",  linha_w, linha_h,
                            pos_x_linha, pos_y, gold_shape, behind=True)
    xml_text  = anchor_xml(11, "SidebarText", box_w, box_h,
                            pos_x_text + linha_w + 30000, pos_y, text_shape)

    for xml in [xml_linha, xml_text]:
        el = etree.fromstring(xml)
        doc.element.body.insert(0, el)


# ═══════════════════════════════════════════════════════════════════════════════
# CORPO DE EXEMPLO
# ═══════════════════════════════════════════════════════════════════════════════

def adicionar_corpo_exemplo(doc: Document):
    """Preenche o body com o conteúdo padrão de um Agravo de Instrumento."""
    safe = lambda nome: nome if nome in [s.name for s in doc.styles] else "Normal"

    def p(txt="", style="AE Corpo", align=None, bold=False, italic=False,
          color=None, size=None, spacing=None, before=None, after=None):
        para = doc.add_paragraph(style=safe(style))
        if align:   para.alignment = align
        if before is not None or after is not None:
            pf = para.paragraph_format
            if before is not None: pf.space_before = Pt(before)
            if after  is not None: pf.space_after  = Pt(after)
        if txt:
            r = para.add_run(txt)
            if bold:    r.font.bold  = True
            if italic:  r.font.italic = True
            if color:   r.font.color.rgb = color
            if size:    r.font.size  = Pt(size)
            if spacing: add_run_spacing(r, spacing)
        return para

    def sep():
        doc.add_paragraph()

    # ── Endereçamento ─────────────────────────────────────────────────────────
    pa = p("AO EXMO. SR. DESEMBARGADOR PRESIDENTE", style="AE Endereçamento")
    p("Tribunal de Justiça do Estado de [UF]",
      style="AE Corpo", italic=True)

    sep()
    p("Autos de origem nº _________________ — ___ª Vara _____ da Comarca de ___________/___",
      style="AE Corpo")
    sep()

    # Qualificação
    p("[NOME DO AGRAVANTE], [nacionalidade], [estado civil], [profissão], portador do RG "
      "nº _______, CPF nº _______________, residente e domiciliado em [endereço completo], "
      "por intermédio de seu advogado que esta subscreve (OAB/SP nº ______), vem, "
      "respeitosamente, à presença de Vossa Excelência, com fundamento nos arts. 1.015, I, "
      "e 1.019, I, do Código de Processo Civil, interpor",
      style="AE Corpo")
    sep()

    # ── Título da peça ────────────────────────────────────────────────────────
    p1 = doc.add_paragraph(style=safe("AE Título 1"))
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("AGRAVO DE INSTRUMENTO")
    r1.font.name = "Georgia"
    r1.font.bold = True
    r1.font.size = Pt(14)
    r1.font.color.rgb = C_NAVY1
    add_run_small_caps(r1)
    add_run_spacing(r1, 60)

    p2 = doc.add_paragraph(style=safe("AE Título 1"))
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("COM PEDIDO DE TUTELA RECURSAL")
    r2.font.name = "Georgia"
    r2.font.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = C_NAVY1
    add_run_small_caps(r2)
    add_run_spacing(r2, 60)

    # Ornamento dourado
    p_orn = p("◆", style="AE Ornamento")
    p_orn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep()

    p("contra decisão proferida pelo juízo a quo nos autos em epígrafe, pelas razões de "
      "fato e de direito a seguir expostas, requerendo seja recebido e processado na "
      "forma da lei, para ao final ser conhecido e provido.",
      style="AE Corpo")
    sep()

    # ── Caixa comparativa ─────────────────────────────────────────────────────
    tab = doc.add_table(1, 2)
    set_table_no_borders(tab)
    _set_table_xml(tab, 9355, 0, [4677, 4678])

    lc = tab.cell(0, 0)
    rc = tab.cell(0, 1)

    # Bordas: todas finas cinza, exceto divisória interna em gold
    for cell in (lc, rc):
        set_cell_border(cell, {
            "top":    ("single", "4", NAVY2),
            "bottom": ("single", "4", NAVY2),
            "left":   ("single", "4", NAVY2),
            "right":  ("single", "4", NAVY2),
        })

    def _cabecalho_tab(cell, titulo):
        p_h = cell.paragraphs[0]
        p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p_h, before_pt=6, after_pt=0, line_rule="auto", line=240)
        r_h = p_h.add_run(titulo)
        r_h.font.name = "Arial"
        r_h.font.size = Pt(8)
        r_h.font.color.rgb = C_NAVY1
        add_run_spacing(r_h, 120)
        add_run_small_caps(r_h)
        # ◆ dourado
        p_o = cell.add_paragraph()
        p_o.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p_o, before_pt=0, after_pt=4, line_rule="auto", line=200)
        r_o = p_o.add_run("◆")
        r_o.font.size = Pt(7)
        r_o.font.color.rgb = C_GOLD

    _cabecalho_tab(lc, "DECISÃO RECORRIDA")
    _cabecalho_tab(rc, "PEDIDO")

    def _corpo_tab(cell, txt):
        p_t = cell.add_paragraph()
        p_t.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_para_spacing(p_t, before_pt=0, after_pt=6, line_rule="auto", line=320)
        r_t = p_t.add_run(txt)
        r_t.font.name = "Times New Roman"
        r_t.font.size = Pt(11)

    _corpo_tab(lc, "[Descrever aqui a decisão que está sendo impugnada, "
                   "indicando o ponto específico objeto da reforma.]")
    _corpo_tab(rc, "[Indicar o pedido de reforma, de forma objetiva e "
                   "diretamente relacionada à decisão recorrida.]")

    sep()

    # ── Justificativa ─────────────────────────────────────────────────────────
    p_just = doc.add_paragraph(style=safe("AE Título 1"))
    p_just.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_just = p_just.add_run("JUSTIFICATIVA DA REFORMA")
    r_just.font.name = "Georgia"
    r_just.font.size = Pt(11)
    r_just.font.color.rgb = C_NAVY1
    add_run_spacing(r_just, 100)
    add_run_small_caps(r_just)

    p_orn2 = p("◆", style="AE Ornamento")
    p_orn2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Itens com marcador dourado
    itens = [
        "[Fundamento legal ou jurisprudencial 1 — descrever de forma concisa "
        "o argumento principal que justifica a reforma da decisão.]",
        "[Fundamento legal ou jurisprudencial 2 — indicar norma, súmula ou "
        "acórdão que ampara o pedido.]",
        "[Fundamento fático 3 — descrever como o caso concreto se enquadra nos "
        "fundamentos jurídicos apresentados.]",
    ]
    for item in itens:
        p_item = doc.add_paragraph()
        p_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_para_spacing(p_item, before_pt=3, after_pt=3, line_rule="auto", line=340)
        pf = p_item.paragraph_format
        pf.left_indent   = Cm(0.8)
        pf.first_line_indent = Cm(-0.8)
        r_bullet = p_item.add_run("◆  ")
        r_bullet.font.size = Pt(8)
        r_bullet.font.color.rgb = C_GOLD
        r_body = p_item.add_run(item)
        r_body.font.name = "Times New Roman"
        r_body.font.size = Pt(12)
    sep()

    # ── Seções da petição ─────────────────────────────────────────────────────
    for sec in [("I — DO CABIMENTO", "O presente recurso é tempestivo e cabível, "
                 "nos termos do art. 1.015, I, do CPC, por versar sobre decisão "
                 "que aprecia tutela provisória."),
                ("II — DA TUTELA RECURSAL", "Presentes a probabilidade do direito "
                 "e o perigo de dano ou risco ao resultado útil do processo "
                 "(art. 300 c/c art. 1.019, I, do CPC), requer-se a concessão "
                 "da tutela recursal inaudita altera pars."),
                ("III — DOS PEDIDOS", "Ante o exposto, requer-se o conhecimento e "
                 "provimento do presente agravo de instrumento, com a reforma "
                 "integral da decisão agravada.")]:

        titulo, corpo = sec
        p_sec = doc.add_paragraph(style=safe("AE Título 2"))
        r_sec = p_sec.add_run(titulo.upper())
        r_sec.font.name = "Georgia"
        r_sec.font.bold = True
        r_sec.font.size = Pt(12)
        r_sec.font.color.rgb = C_NAVY1
        add_run_spacing(r_sec, 30)

        p(corpo, style="AE Corpo")
        sep()

    # ── Assinatura ────────────────────────────────────────────────────────────
    p("Nestes termos, pede deferimento.", style="AE Corpo")
    sep()
    p(f"[Cidade], _____ de _____________ de _______.", style="AE Corpo")
    sep(); sep()

    p_ass = p("", style="AE Corpo", align=WD_ALIGN_PARAGRAPH.CENTER)
    p_ass.clear()
    p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ln = p_ass.add_run("_" * 42)
    r_ln.font.color.rgb = C_NAVY1

    p_nome = p("Vanildo Angelo Epifanio", style="AE Corpo",
               align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=C_NAVY1)
    p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_oab = p(f"Advogado — {OAB}", style="AE Nota", align=WD_ALIGN_PARAGRAPH.CENTER)
    p_oab.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def criar_template(tipo_peca: str = "AGRAVO DE INSTRUMENTO"):
    doc = Document()

    # Margens
    section = doc.sections[0]
    section.page_height          = Cm(29.7)
    section.page_width           = Cm(21.0)
    section.top_margin           = Cm(2.0)
    section.bottom_margin        = Cm(2.0)
    section.left_margin          = Cm(2.5)
    section.right_margin         = Cm(2.0)
    section.header_distance      = Cm(0.5)
    section.footer_distance      = Cm(0.0)
    section.different_first_page_header_footer = False

    # Estilos
    criar_estilos(doc)

    # Cabeçalho e rodapé
    criar_cabecalho(section)
    criar_rodape(section)

    # Corpo de exemplo
    adicionar_corpo_exemplo(doc)

    # Salva primeiro (o sidebar precisa reabrir o doc)
    out = MODELOS / "AGRAVO_INSTRUMENTO.docx"
    doc.save(str(out))

    # Adiciona sidebar
    doc2 = Document(str(out))
    adicionar_sidebar(doc2, tipo_peca)
    doc2.save(str(out))

    print(f"Template premium salvo: {out}")
    return out


if __name__ == "__main__":
    criar_template()
